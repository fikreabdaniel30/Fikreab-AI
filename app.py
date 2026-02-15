import streamlit as st
import google.generativeai as genai
from PyPDF2 import PdfReader
from docx import Document
from fpdf import FPDF
from io import BytesIO

# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="Fikreab AI | Ultimate Study Companion",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ---------------- CUSTOM CSS ----------------
st.markdown("""
<style>
.stApp {
    background: radial-gradient(circle at top right, #1e2127, #0e1117);
    color: #ffffff;
}

.header-banner {
    background: linear-gradient(90deg, #00C9FF 0%, #92FE9D 100%);
    padding: 30px;
    border-radius: 20px;
    text-align: center;
    margin-bottom: 30px;
    box-shadow: 0 10px 30px rgba(0, 201, 255, 0.2);
}

.header-banner h1 {
    color: #000 !important;
    margin: 0;
    font-weight: 800;
    font-size: 3rem;
}

.css-card {
    border-radius: 15px;
    padding: 25px;
    background-color: rgba(255, 255, 255, 0.05);
    backdrop-filter: blur(12px);
    border: 1px solid rgba(255, 255, 255, 0.1);
    margin-bottom: 20px;
}

.stButton>button {
    width: 100%;
    border-radius: 12px;
    height: 3.5em;
    background: linear-gradient(90deg, #00C9FF 0%, #92FE9D 100%);
    color: #000 !important;
    font-weight: bold;
    border: none;
}

[data-testid="stSidebar"] {
    background-color: #0a0c10;
}
</style>
""", unsafe_allow_html=True)

# ---------------- EXPORT FUNCTIONS ----------------
def get_docx(text):
    doc = Document()
    doc.add_heading('Fikreab AI Study Notes', 0)

    clean_text = text.replace('**', '').replace('#', '')

    for paragraph in clean_text.split('\n'):
        if paragraph.strip():
            doc.add_paragraph(paragraph)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def get_pdf(text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    safe_text = "".join(i for i in text if ord(i) < 128)
    clean_text = safe_text.replace('**', '').replace('#', '')

    for line in clean_text.split('\n'):
        pdf.multi_cell(0, 10, txt=line)

    return pdf.output(dest='S').encode('latin-1', 'ignore')


# ---------------- GEMINI SETUP ----------------
model = None

if "GEMINI_API_KEY" in st.secrets:
    try:
        api_key = st.secrets["GEMINI_API_KEY"].strip()
        genai.configure(AIzaSyBn6UPO8tQ-c0YNUbW7oAG9WGoO2Xfs5_Y)

        # Auto detect available models
        available_models = [
            m.name for m in genai.list_models()
            if "generateContent" in m.supported_actions
        ]

        preferred = [
            "models/gemini-1.5-flash",
            "models/gemini-2.0-flash",
            "models/gemini-pro"
        ]

        selected_model = next(
            (m for m in preferred if m in available_models),
            available_models[0]
        )

        model = genai.GenerativeModel(selected_model)

    except Exception as e:
        st.error(f"❌ Gemini Setup Error: {e}")

else:
    st.sidebar.error("❌ GEMINI_API_KEY missing in Streamlit secrets")


# ---------------- SIDEBAR ----------------
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/4712/4712035.png", width=80)
    st.title("Fikreab AI")

    st.markdown("---")

    ai_mode = st.selectbox(
        "Select Action:",
        [
            "📝 Comprehensive Notes",
            "🗂️ Flashcard Set",
            "📉 Key Points Only",
            "🎯 Exam Predictions"
        ]
    )

    uploaded_file = st.file_uploader("Upload PDF", type="pdf")

    if uploaded_file:
        st.success("✅ Document Loaded")


# ---------------- HEADER ----------------
st.markdown(
    '<div class="header-banner"><h1>🎓 FIKREAB AI STUDIO</h1>'
    '<p style="color:black;">Advanced Academic Intelligence</p></div>',
    unsafe_allow_html=True
)

# ---------------- FILE PROCESSING ----------------
text_content = ""

if uploaded_file:
    reader = PdfReader(uploaded_file)

    text_content = "".join(
        [page.extract_text() or "" for page in reader.pages]
    )[:30000]

# ---------------- MAIN UI ----------------
if uploaded_file:

    tab1, tab2 = st.tabs(["🚀 Study Generator", "🧠 Quiz Engine"])

    # ---------------- GENERATOR ----------------
    with tab1:

        col1, col2 = st.columns([1, 2])

        with col1:
            st.markdown(f"### 📍 Mode\n**{ai_mode}**")
            generate_btn = st.button("✨ START GENERATION")

        with col2:

            if generate_btn:

                if not model:
                    st.error("Model not initialized")
                else:

                    prompts = {
                        "📝 Comprehensive Notes":
                            "Create detailed structured study notes with headings and explanations.",

                        "🗂️ Flashcard Set":
                            "Create flashcards in format Front: ... Back: ...",

                        "📉 Key Points Only":
                            "Summarize into high impact bullet points.",

                        "🎯 Exam Predictions":
                            "Predict 5 exam questions with answers."
                    }

                    with st.spinner("⏳ Generating..."):
                        try:
                            response = model.generate_content(
                                f"{prompts[ai_mode]}\n\nContent:\n{text_content}"
                            )

                            st.session_state["output"] = response.text
                            st.toast("✅ Generation Complete")

                        except Exception as e:
                            st.error(f"API Error: {e}")

            if "output" in st.session_state:
                st.markdown("<div class='css-card'>", unsafe_allow_html=True)
                st.markdown(st.session_state["output"])
                st.markdown("</div>", unsafe_allow_html=True)

                st.write("### 📥 Export")

                c1, c2 = st.columns(2)

                with c1:
                    st.download_button(
                        "📄 Export PDF",
                        get_pdf(st.session_state["output"]),
                        "Fikreab_AI.pdf"
                    )

                with c2:
                    st.download_button(
                        "📝 Export Word",
                        get_docx(st.session_state["output"]),
                        "Fikreab_AI.docx"
                    )

    # ---------------- QUIZ ENGINE ----------------
    with tab2:

        if st.button("🎯 Generate Quiz"):

            if model:
                with st.spinner("Creating quiz..."):
                    try:
                        resp = model.generate_content(
                            f"Create 5 MCQ questions with answers:\n{text_content}"
                        )

                        st.session_state["quiz"] = resp.text

                    except Exception as e:
                        st.error(f"Quiz Error: {e}")

        if "quiz" in st.session_state:
            st.markdown(
                f"<div class='css-card'>{st.session_state['quiz']}</div>",
                unsafe_allow_html=True
            )

# ---------------- EMPTY STATE ----------------
else:

    st.markdown("""
    <div style='text-align: center; padding: 60px;'>
        <h2 style='color: #00e5ff;'>Ready to upgrade your grades?</h2>
        <p>Upload your PDF in the sidebar to begin.</p>
        <div style='font-size: 5rem;'>📚</div>
    </div>
    """, unsafe_allow_html=True)

